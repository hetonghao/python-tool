"""
@author HeTongHao
@date 2019/4/1 11:00
@description  读取所有数据表,导出Word文件
"""
from export_table.export_pgsql import list_group_result, list_column_detail, value_handle
from docx import Document
from docx.shared import Inches


class ExportWord:

    def __init__(self, export_path='export_dir/'):
        self.export_path = export_path

    def export(self):
        module_greps = list_group_result()  # 分组后的结果
        for grep_name in module_greps.keys():
            doc = get_default_document()  # 获取默认样式的文档
            for table_detail in module_greps[grep_name]:
                add_table(doc, table_detail)
            file_name = 'bull_' + grep_name + '模块表结构设计.docx'
            doc.save(self.export_path + file_name)
            print(file_name + '导出成功!')
        print('-----------所有数据表结构导出完毕!')


def get_default_document():
    document = Document()
    distance = Inches(0.3)
    sec = document.sections[0]  # sections对应文档中的“节”
    sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance
    sec.page_width = Inches(12)  # 设置页面宽度
    sec.page_height = Inches(20)  # 设置页面高度
    return document


def add_table(doc, pg_table_detail):
    pg_table_name = pg_table_detail[0]
    pg_table_dec = pg_table_detail[1] if pg_table_detail[1] is not None else ''
    columns = list_column_detail(pg_table_name)
    doc.add_heading(pg_table_name + '  ' + pg_table_dec)
    column_number = str(columns[0]).replace('(', '').replace(')', '').split(',').__len__()
    table = doc.add_table(len(columns) + 1, column_number, 'Light Shading Accent 1')  # 添加一个空表
    first_row = table.rows[0]
    first_row.cells[0].text = '字段名'
    first_row.cells[1].text = '类型'
    first_row.cells[2].text = '是否为空'
    first_row.cells[3].text = '描述'
    for column, row in zip(columns, table.rows[1:]):
        column_values = str(column).replace('(', '').replace(')', '').split(',')
        for column_value, cell in zip(column_values, row.cells):
            value = column_value.replace('\'', '').strip()
            value = value_handle(value)
            cell.text = value
